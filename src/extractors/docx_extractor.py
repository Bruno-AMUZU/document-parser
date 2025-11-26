"""Extracteur pour les fichiers DOCX."""

from docx import Document as DocxDocument
from pathlib import Path
from src.extractors.base_extractor import BaseExtractor
from src.models.document import Document


class DocxExtractor(BaseExtractor):
    """Extracteur pour les fichiers DOCX (.docx)."""
    
    def extract(self, file_path: str) -> Document:
        """
        Extrait le contenu d'un fichier DOCX.
        
        Args:
            file_path: Chemin vers le fichier DOCX
            
        Returns:
            Document: Objet Document contenant le contenu extrait
            
        Raises:
            FileNotFoundError: Si le fichier n'existe pas
            Exception: Si le DOCX ne peut pas être lu
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")
        
        content_parts = []
        metadata = {}
        
        try:
            doc = DocxDocument(file_path)
            
            # Extraction du texte de tous les paragraphes
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # Extraction du texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        content_parts.append(row_text)
            
            # Extraction des propriétés du document
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'category': core_props.category,
                'comments': core_props.comments,
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'revision': core_props.revision,
                'last_modified_by': core_props.last_modified_by
            }
            
            # Nettoyage des valeurs None
            metadata = {k: v for k, v in metadata.items() if v is not None}
        
        except Exception as e:
            raise Exception(f"Erreur lors de la lecture du DOCX: {str(e)}")
        
        content = '\n\n'.join(content_parts)
        
        return Document(
            content=content,
            file_path=str(path.absolute()),
            file_type=self.get_file_type(file_path),
            title=metadata.get('title'),
            author=metadata.get('author'),
            metadata=metadata
        )

